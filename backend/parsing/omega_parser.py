# Объединение парсеров в один.

# TODO: .doc

import collections
from io import StringIO
import json
import os
from pathlib import Path
import re
import xml
import xml.etree
import xml.etree.ElementTree

from bs4 import BeautifulSoup as bs
import docx  # python-docx
import docx2txt2
from docx_parser import DocumentParser  # docx_parser
from markdown import Markdown
from striprtf.striprtf import rtf_to_text  # striprtf
import pptx2txt2
from pydocx import PyDocX

from parsing.pdf_parser import pdf_to_string, clean_text



collections.Hashable = collections.abc.Hashable


class MarkdownToPlain:
    def __init__(self):
        Markdown.output_formats["plain"] = MarkdownToPlain.unmark_element
        self.md = Markdown(output_format="plain")
        self.md.stripTopLevelTags = False

    @staticmethod
    def unmark_element(element: xml.etree.ElementTree.Element, stream: StringIO | None = None) -> str:
        if stream is None:
            stream = StringIO()
        if element.text:
            stream.write(element.text)
        for sub in element:
            MarkdownToPlain.unmark_element(sub, stream)
        if element.tail:
            stream.write(element.tail)
        return stream.getvalue()

    def convert(self, source: str) -> str:
        return self.md.convert(source)


md = MarkdownToPlain()


def extract_text_from_markdown(path: os.PathLike) -> str:
    with open(path) as f:
        return md.convert(f.read())


def extract_text_from_rtf(path: os.PathLike) -> str:
    with open(path) as f:
        text = rtf_to_text(f.read())
        text = re.sub(r"\|", " ", text)
        return text


def extract_text_from_docx(path: os.PathLike) -> str:
    """
    Извлекает текст из файла формата DOCX (MS OFFICE 2007+).
    Зачастую бросает ошибку, связанную со структурой документа.

    Parameters:
    - path (os.PathLike): Путь к файлу.

    Returns:
    - str: текст элемента
    """
    res: list[str] = []
    doc = DocumentParser(path)
    for _type, item in doc.parse():
        if "text" in item:
            res.append(item["text"])
        elif "data" in item:
            res.extend(" ".join(row) for row in item["data"])

    return "\n".join(res)


def extract_text_from_docx_2(path: os.PathLike) -> str:
    document = docx.Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(" ".join(c.text() for c in row.cells)
                       for table in document.tables for row in table.rows)

    return f"{text}\n{tables}"


def extract_text_from_doc(path: os.PathLike) -> str:
    html = PyDocX.to_html(path)
    return extract_text_with_bs(html)


def extract_text_with_bs(source: str) -> str:
    soup = bs(source)
    [s.extract() for s in soup(["style", "script"])]
    tmpText = soup.get_text(separator="\n")
    return tmpText


def read_any_doc(path: os.PathLike) -> str:
    text: str = ""
    extention: str = path.suffix

    if extention == ".txt":
        with open(path) as f:
            text = f.read()
    elif extention == ".md":
        text = extract_text_from_markdown(path)
    elif extention == ".pdf":
        text = pdf_to_string(path)
    elif extention == ".rtf":
        text = extract_text_from_rtf(path)
    elif extention == ".pptx" or extention == ".odp":
        text = pptx2txt2.extract_text(path)
    elif extention == ".odt":
        text = docx2txt2.extract_text(path)
    elif extention == ".docx":
        try:
            text = extract_text_from_docx(path)
        except Exception as e:
            try:
                text = extract_text_from_docx_2(path)
            except Exception as e:
                # могут присутствовать артефакты
                text = docx2txt2.extract_text(path)
    elif extention == ".doc":
        text = extract_text_from_doc(path)

        if text == '':
            text = docx2txt2.extract_text(path)
    else:
        with open(path, errors="ignore") as f:
            text = extract_text_with_bs(f.read())
    return clean_text(text)

